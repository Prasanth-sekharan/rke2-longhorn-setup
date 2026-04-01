#!/usr/bin/env python3
"""
Generate professional RKE2 + Longhorn deployment document in .docx format.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────────────────
# PLACEHOLDER VALUES — Replace these before finalizing
# ──────────────────────────────────────────────────────────
PLACEHOLDERS = {
    "DOMAIN": "<DOMAIN_NAME>",
    "K8S_VERSION": "<KUBERNETES_VERSION>",
    "RKE2_VERSION": "<RKE2_VERSION>",
    "LONGHORN_VERSION": "<LONGHORN_VERSION>",
    "ORG_NAME": "<ORGANIZATION_NAME>",
    "PROJECT_NAME": "<PROJECT_NAME>",
    "ENVIRONMENT": "<ENVIRONMENT>",  # e.g., Production, Staging
    "DATE": datetime.date.today().strftime("%B %d, %Y"),
    "AUTHOR": "<AUTHOR_NAME>",
    "APPROVER": "<APPROVER_NAME>",
    "DOCUMENT_ID": "<DOCUMENT_ID>",
    "SERVER1_NAME": "<SERVER1_HOSTNAME>",
    "SERVER1_IP": "<SERVER1_IP>",
    "SERVER2_NAME": "<SERVER2_HOSTNAME>",
    "SERVER2_IP": "<SERVER2_IP>",
    "SERVER3_NAME": "<SERVER3_HOSTNAME>",
    "SERVER3_IP": "<SERVER3_IP>",
    "AGENT1_NAME": "<AGENT1_HOSTNAME>",
    "AGENT1_IP": "<AGENT1_IP>",
    "AGENT2_NAME": "<AGENT2_HOSTNAME>",
    "AGENT2_IP": "<AGENT2_IP>",
    "LB_VIP": "<LOAD_BALANCER_VIP>",
    "LB_FQDN": "<LOAD_BALANCER_FQDN>",
    "NETWORK_CIDR": "<POD_NETWORK_CIDR>",
    "SVC_CIDR": "<SERVICE_CIDR>",
    "NTP_SERVER": "<NTP_SERVER>",
    "DNS_SERVER": "<DNS_SERVER>",
    "OS_VERSION": "<OS_VERSION>",
    "S3_ENDPOINT": "<S3_ENDPOINT>",
    "S3_BUCKET": "<S3_BUCKET>",
    "S3_REGION": "<S3_REGION>",
}

P = PLACEHOLDERS  # shorthand


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F4F7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_code_block(doc, text, font_size=8):
    """Add a code block with monospace font and background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(33, 33, 33)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F0F0F0")
    shading.set(qn("w:val"), "clear")
    p._element.get_or_add_pPr().append(shading)

    return p


def add_note(doc, text, label="NOTE"):
    """Add a note/callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(f"[{label}] ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = RGBColor(180, 80, 0)
    run_text = p.add_run(text)
    run_text.font.size = Pt(9)
    run_text.font.italic = True


def build_document():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(P["ORG_NAME"])
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(P["PROJECT_NAME"])
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RKE2 Kubernetes Cluster\nwith Longhorn Distributed Storage")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(30, 50, 80)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Deployment & Operations Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(4):
        doc.add_paragraph()

    # Document metadata table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document ID", P["DOCUMENT_ID"]),
        ("Version", "1.0"),
        ("Date", P["DATE"]),
        ("Author", P["AUTHOR"]),
        ("Approved By", P["APPROVER"]),
        ("Classification", "Internal"),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # REVISION HISTORY
    # ════════════════════════════════════════════════════════
    doc.add_heading("Revision History", level=1)
    add_table_with_header(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["0.1", "<DATE>", P["AUTHOR"], "Initial draft"],
            ["0.2", "<DATE>", P["AUTHOR"], "Technical review updates"],
            ["1.0", P["DATE"], P["AUTHOR"], "Final release"],
        ],
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # TABLE OF CONTENTS PLACEHOLDER
    # ════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.add_run("[Insert Table of Contents — In Word: References → Table of Contents]").italic = True
    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        f"This document provides the technical specification and deployment procedures for a "
        f"high-availability Kubernetes cluster using RKE2 {P['RKE2_VERSION']} with Longhorn "
        f"{P['LONGHORN_VERSION']} distributed block storage. The cluster is deployed across "
        f"3 server nodes (control plane + etcd + worker) and 2 agent nodes (worker) within "
        f"the {P['ENVIRONMENT']} environment of {P['ORG_NAME']}."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph("This document covers:")
    for item in [
        "Cluster architecture and topology design",
        "Infrastructure prerequisites and network configuration",
        "RKE2 installation and cluster bootstrap procedures",
        "Longhorn deployment and storage configuration",
        "PersistentVolume and StorageClass provisioning",
        "Operational procedures and maintenance tasks",
        "Scaling procedures (horizontal and vertical)",
        "Backup, restore, and disaster recovery",
        "Troubleshooting and diagnostic procedures",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.3 Document Conventions", level=2)
    doc.add_paragraph(
        "Throughout this document, placeholder values are enclosed in angle brackets "
        "and must be substituted with environment-specific values prior to execution. "
        "All placeholders are defined in Section 1.4."
    )

    doc.add_heading("1.4 Placeholder Reference", level=2)
    add_table_with_header(
        doc,
        ["Placeholder", "Description", "Example Value"],
        [
            ["<DOMAIN_NAME>", "Internal DNS domain", "corp.example.com"],
            ["<KUBERNETES_VERSION>", "Kubernetes version", "v1.29.4"],
            ["<RKE2_VERSION>", "RKE2 release version", "v1.29.4+rke2r1"],
            ["<LONGHORN_VERSION>", "Longhorn chart version", "v1.6.2"],
            ["<SERVER{n}_HOSTNAME>", "Control plane node FQDN", "rke2-cp-01.corp.example.com"],
            ["<SERVER{n}_IP>", "Control plane node IP", "10.0.1.11"],
            ["<AGENT{n}_HOSTNAME>", "Worker node FQDN", "rke2-wk-01.corp.example.com"],
            ["<AGENT{n}_IP>", "Worker node IP", "10.0.1.21"],
            ["<LOAD_BALANCER_VIP>", "API server VIP/endpoint", "10.0.1.100"],
            ["<LOAD_BALANCER_FQDN>", "API server FQDN", "k8s-api.corp.example.com"],
            ["<POD_NETWORK_CIDR>", "Pod network CIDR", "10.42.0.0/16"],
            ["<SERVICE_CIDR>", "Service cluster CIDR", "10.43.0.0/16"],
            ["<NTP_SERVER>", "NTP synchronization source", "ntp.corp.example.com"],
            ["<DNS_SERVER>", "DNS resolver address", "10.0.0.2"],
            ["<S3_ENDPOINT>", "S3 backup endpoint", "s3.amazonaws.com"],
            ["<S3_BUCKET>", "S3 backup bucket", "rke2-backups"],
        ],
        col_widths=[5, 7, 6],
    )

    doc.add_heading("1.5 Reference Documents", level=2)
    for item in [
        "RKE2 Official Documentation — https://docs.rke2.io",
        "Longhorn Official Documentation — https://longhorn.io/docs",
        "Kubernetes Documentation — https://kubernetes.io/docs",
        "CIS Kubernetes Benchmark v1.8",
        f"{P['ORG_NAME']} Infrastructure Standards",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ════════════════════════════════════════════════════════
    doc.add_heading("2. Cluster Architecture", level=1)

    doc.add_heading("2.1 Topology Overview", level=2)
    doc.add_paragraph(
        "The cluster comprises 5 nodes deployed in a 3+2 topology: 3 server nodes "
        "providing high-availability control plane with embedded etcd, and 2 agent nodes "
        "providing dedicated compute capacity. All nodes are schedulable for workload "
        "execution, including the control plane nodes."
    )

    add_table_with_header(
        doc,
        ["Node", "Hostname", "IP Address", "Role", "CPU", "RAM", "Disk"],
        [
            ["SERVER-1", P["SERVER1_NAME"], P["SERVER1_IP"], "Control Plane + etcd + Worker", "4+", "8GB+", "2TB"],
            ["SERVER-2", P["SERVER2_NAME"], P["SERVER2_IP"], "Control Plane + etcd + Worker", "4+", "8GB+", "2TB"],
            ["SERVER-3", P["SERVER3_NAME"], P["SERVER3_IP"], "Control Plane + etcd + Worker", "4+", "8GB+", "2TB"],
            ["AGENT-1", P["AGENT1_NAME"], P["AGENT1_IP"], "Worker", "4+", "8GB+", "50GB+"],
            ["AGENT-2", P["AGENT2_NAME"], P["AGENT2_IP"], "Worker", "4+", "8GB+", "50GB+"],
        ],
    )

    doc.add_heading("2.2 Control Plane Components", level=2)
    doc.add_paragraph(
        "Each server node runs the complete Kubernetes control plane stack. RKE2 deploys "
        "these components as static pods managed by kubelet."
    )
    add_table_with_header(
        doc,
        ["Component", "Function", "Port(s)", "Protocol"],
        [
            ["kube-apiserver", "REST API endpoint for cluster management", "6443", "HTTPS"],
            ["kube-scheduler", "Pod placement and node assignment", "-", "gRPC"],
            ["kube-controller-manager", "Control loop reconciliation", "-", "gRPC"],
            ["etcd", "Distributed key-value store (cluster state)", "2379, 2380", "HTTPS"],
            ["kubelet", "Node agent, static pod management", "10250", "HTTPS"],
            ["kube-proxy", "Service networking (iptables/IPVS)", "-", "gRPC"],
            ["cloud-controller-manager", "Cloud provider integration (if applicable)", "-", "gRPC"],
        ],
    )

    doc.add_heading("2.3 etcd Cluster Configuration", level=2)
    doc.add_paragraph(
        "The etcd cluster runs as a 3-member embedded etcd deployment within the RKE2 "
        "server nodes. This configuration provides fault tolerance for a single node failure."
    )
    add_table_with_header(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["Member count", "3", "Odd number for quorum"],
            ["Quorum", "2/3", "Minimum nodes for consensus"],
            ["Fault tolerance", "1 node", "Survives 1 simultaneous failure"],
            ["Backend", "bbolt (embedded)", "Default RKE2 etcd backend"],
            ["Snapshot interval", "6 hours", "Automated snapshots"],
            ["Snapshot retention", "5 snapshots", "Oldest snapshots pruned"],
            ["Default data dir", "/var/lib/rancher/rke2/server/db", "etcd data storage"],
        ],
    )

    doc.add_heading("2.4 Network Architecture", level=2)
    doc.add_paragraph(
        "RKE2 deploys Canal (Calico + Flannel) as the default CNI, providing pod-to-pod "
        "networking via VXLAN overlay and network policy enforcement via Calico."
    )
    add_table_with_header(
        doc,
        ["Network", "CIDR", "Purpose"],
        [
            ["Pod Network", P["NETWORK_CIDR"], "Pod IP allocation (Calico IPAM)"],
            ["Service Network", P["SVC_CIDR"], "ClusterIP service allocation"],
            ["Node Network", P["NETWORK_CIDR"].rsplit(".", 1)[0] + ".0/24 (example)", "Host network (physical/VLAN)"],
            ["VXLAN Overlay", "flannel.1 interface", "Cross-node pod-to-pod traffic"],
        ],
    )

    doc.add_heading("2.5 High Availability Analysis", level=2)
    add_table_with_header(
        doc,
        ["Component", "Failure Mode", "Impact", "Recovery"],
        [
            ["etcd (1 node down)", "Quorum maintained (2/3)", "No impact", "Auto-recovery on restart"],
            ["etcd (2 nodes down)", "Quorum lost", "Cluster read-only", "Restore from snapshot"],
            ["API Server", "Load balancer routes to healthy nodes", "No impact", "Auto-recovery"],
            ["Longhorn (1 replica)", "2/3 replicas active", "Degraded volume", "Auto-rebuild to healthy node"],
            ["Longhorn (2 replicas)", "1/3 replica active", "Data accessible, no redundancy", "Manual intervention required"],
            ["Agent node", "Pods rescheduled", "Brief service interruption", "Kubernetes self-healing"],
        ],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 3. PREREQUISITES
    # ════════════════════════════════════════════════════════
    doc.add_heading("3. Infrastructure Prerequisites", level=1)

    doc.add_heading("3.1 Operating System Requirements", level=2)
    add_table_with_header(
        doc,
        ["Requirement", "Specification"],
        [
            ["Supported OS", f"{P['OS_VERSION']} (RHEL 8/9, Rocky Linux 8/9, Ubuntu 20.04/22.04)"],
            ["Kernel", "≥ 4.18 (RHEL 8+), ≥ 5.4 (Ubuntu 20.04+)"],
            ["SELinux", "Disabled or Permissive"],
            ["Firewall", "Disabled (RKE2 manages iptables)"],
            ["Swap", "Disabled"],
            ["Time Sync", f"NTP configured, pointing to {P['NTP_SERVER']}"],
            ["DNS Resolution", f"Forward/reverse DNS configured, resolver: {P['DNS_SERVER']}"],
            ["Filesystem", "XFS or ext4 on data partitions"],
            ["cgroup v2", "Supported (RKE2 auto-detects)"],
        ],
    )

    doc.add_heading("3.2 Network Port Requirements", level=2)
    doc.add_paragraph("The following ports must be open between all cluster nodes:")
    add_table_with_header(
        doc,
        ["Port", "Direction", "Protocol", "Component", "Nodes"],
        [
            ["6443", "Inbound", "TCP", "Kubernetes API", "All"],
            ["9345", "Inbound", "TCP", "RKE2 node registration", "Server nodes"],
            ["2379", "Inbound", "TCP", "etcd client", "Server nodes"],
            ["2380", "Inbound", "TCP", "etcd peer", "Server nodes"],
            ["10250", "Inbound", "TCP", "kubelet API", "All"],
            ["4789", "Bidirectional", "UDP", "VXLAN (Flannel/Canal)", "All"],
            ["8472", "Bidirectional", "UDP", "VXLAN (Flannel/Canal)", "All"],
            ["30000-32767", "Inbound", "TCP/UDP", "NodePort services", "All"],
            ["9500-9502", "Inbound", "TCP", "Longhorn manager/data", "All"],
            ["31080", "Inbound", "TCP", "Longhorn UI (NodePort)", "All"],
        ],
    )

    doc.add_heading("3.3 DNS Configuration", level=2)
    doc.add_paragraph(
        "Forward and reverse DNS records must be configured for all nodes. "
        "The API server endpoint must resolve to the load balancer VIP."
    )
    add_table_with_header(
        doc,
        ["Record Type", "Name", "Value"],
        [
            ["A", P["SERVER1_NAME"], P["SERVER1_IP"]],
            ["A", P["SERVER2_NAME"], P["SERVER2_IP"]],
            ["A", P["SERVER3_NAME"], P["SERVER3_IP"]],
            ["A", P["AGENT1_NAME"], P["AGENT1_IP"]],
            ["A", P["AGENT2_NAME"], P["AGENT2_IP"]],
            ["A", P["LB_FQDN"], P["LB_VIP"]],
            ["PTR", P["SERVER1_IP"], P["SERVER1_NAME"]],
            ["PTR", P["SERVER2_IP"], P["SERVER2_NAME"]],
            ["PTR", P["SERVER3_IP"], P["SERVER3_NAME"]],
        ],
    )

    doc.add_heading("3.4 Load Balancer Configuration", level=2)
    doc.add_paragraph(
        f"A Layer 4 (TCP) load balancer or DNS round-robin must be configured to front "
        f"the Kubernetes API server on port 6443, distributing traffic across the three "
        f"server nodes."
    )
    add_table_with_header(
        doc,
        ["Parameter", "Value"],
        [
            ["VIP / FQDN", f"{P['LB_FQDN']} ({P['LB_VIP']})"],
            ["Frontend Port", "6443"],
            ["Backend Servers", f"{P['SERVER1_IP']}:6443, {P['SERVER2_IP']}:6443, {P['SERVER3_IP']}:6443"],
            ["Health Check", "TCP check on port 6443, interval 5s"],
            ["Algorithm", "Round-robin or least-connections"],
            ["Sticky Sessions", "Not required"],
        ],
    )

    doc.add_heading("3.5 Pre-Installation Validation", level=2)
    doc.add_paragraph("Execute the following validation commands on all nodes prior to installation:")
    add_code_block(
        doc,
        f"# Verify hostname resolution\n"
        f"hostnamectl set-hostname <NODE_FQDN>\n"
        f"ping -c 3 {P['SERVER1_NAME']}\n"
        f"ping -c 3 {P['SERVER2_NAME']}\n"
        f"ping -c 3 {P['SERVER3_NAME']}\n\n"
        f"# Verify NTP synchronization\n"
        f"chronyc sources -v\n"
        f"timedatectl status\n\n"
        f"# Verify swap is disabled\n"
        f"free -h  # Swap line should show 0\n"
        f"swapoff -a\n"
        f"sed -i '/ swap / s/^/#/' /etc/fstab\n\n"
        f"# Verify SELinux status\n"
        f"getenforce  # Should return Disabled or Permissive\n\n"
        f"# Verify firewall is stopped\n"
        f"systemctl is-active firewalld  # Should return inactive\n"
        f"systemctl stop firewalld && systemctl disable firewalld\n\n"
        f"# Verify required ports are not in use\n"
        f"ss -tlnp | grep -E '6443|9345|2379|2380|10250'\n\n"
        f"# Verify disk layout (server nodes)\n"
        f"lsblk\n"
        f"df -hT"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 4. RKE2 INSTALLATION
    # ════════════════════════════════════════════════════════
    doc.add_heading("4. RKE2 Installation Procedures", level=1)

    doc.add_heading("4.1 Deployment Sequence", level=2)
    doc.add_paragraph(
        "Nodes must be deployed in the following order. Each subsequent step depends "
        "on the successful completion of the previous step."
    )
    add_table_with_header(
        doc,
        ["Step", "Node", "Action", "Dependency"],
        [
            ["1", P["SERVER1_NAME"], "Bootstrap first server node", "None"],
            ["2", P["SERVER2_NAME"], "Join second server node", "Step 1 token"],
            ["3", P["SERVER3_NAME"], "Join third server node", "Step 1 token"],
            ["4", P["AGENT1_NAME"], "Join first agent node", "Step 1 token"],
            ["5", P["AGENT2_NAME"], "Join second agent node", "Step 1 token"],
            ["6", "Any server node", "Deploy Longhorn storage", "Steps 1-5"],
            ["7", "Any server node", "Apply manifests", "Step 6"],
        ],
    )

    doc.add_heading("4.2 Server Node Bootstrap (SERVER-1)", level=2)
    doc.add_paragraph(
        f"Execute on {P['SERVER1_NAME']} ({P['SERVER1_IP']}). This node initializes the "
        f"cluster and generates the join token required by subsequent nodes."
    )

    add_code_block(
        doc,
        f"# SSH to the bootstrap node\n"
        f"ssh root@{P['SERVER1_IP']}\n\n"
        f"# Run the bootstrap installation script\n"
        f"sudo bash install-server.sh bootstrap"
    )

    add_note(
        doc,
        "The bootstrap process generates a node-token at "
        "/var/lib/rancher/rke2/server/node-token. Record this token — it is "
        "required for all subsequent node joins.",
        "CRITICAL"
    )

    doc.add_heading("4.2.1 Server Configuration (config.yaml)", level=2)
    doc.add_paragraph("The following configuration is applied on each server node:")
    add_code_block(
        doc,
        f"# /etc/rancher/rke2/config.yaml\n\n"
        f"write-kubeconfig-mode: \"0644\"\n"
        f"tls-san:\n"
        f"  - \"{P['SERVER1_NAME']}\"\n"
        f"  - \"{P['SERVER2_NAME']}\"\n"
        f"  - \"{P['SERVER3_NAME']}\"\n"
        f"  - \"{P['LB_FQDN']}\"\n"
        f"  - \"{P['LB_VIP']}\"\n"
        f"  - \"127.0.0.1\"\n"
        f"  - \"kubernetes\"\n"
        f"  - \"kubernetes.default\"\n"
        f"  - \"kubernetes.default.svc\"\n"
        f"  - \"kubernetes.default.svc.cluster.local\"\n\n"
        f"# Allow workload scheduling on control plane nodes\n"
        f"node-taint: []\n\n"
        f"# etcd snapshot configuration\n"
        f"etcd-snapshot-schedule-cron: \"0 */6 * * *\"\n"
        f"etcd-snapshot-retention: 5\n\n"
        f"# Data directory\n"
        f"data-dir: \"/var/lib/rancher/rke2\"\n\n"
        f"# Network configuration\n"
        f"cluster-cidr: \"{P['NETWORK_CIDR']}\"\n"
        f"service-cidr: \"{P['SVC_CIDR']}\"\n\n"
        f"# CNI (default: canal)\n"
        f"# cni: canal\n\n"
        f"# kubelet arguments\n"
        f"kubelet-arg:\n"
        f"  - \"max-pods=110\"\n"
        f"  - \"image-gc-high-threshold=85\"\n"
        f"  - \"image-gc-low-threshold=80\"\n\n"
        f"# kube-apiserver arguments\n"
        f"kube-apiserver-arg:\n"
        f"  - \"anonymous-auth=false\"\n"
        f"  - \"audit-log-path=/var/lib/rancher/rke2/server/logs/audit.log\"\n"
        f"  - \"audit-log-maxage=30\"\n"
        f"  - \"audit-log-maxbackup=10\"\n"
        f"  - \"audit-log-maxsize=100\""
    )

    doc.add_heading("4.3 Server Node Join (SERVER-2, SERVER-3)", level=2)
    doc.add_paragraph(
        f"Execute on {P['SERVER2_NAME']} and {P['SERVER3_NAME']}. These nodes join the "
        f"existing cluster using the token generated by SERVER-1."
    )
    add_code_block(
        doc,
        f"# SSH to each server node\n"
        f"ssh root@{P['SERVER2_IP']}\n\n"
        f"# Run the join script with SERVER-1 IP\n"
        f"sudo bash install-server.sh join {P['SERVER1_IP']}\n\n"
        f"# Repeat for SERVER-3\n"
        f"ssh root@{P['SERVER3_IP']}\n"
        f"sudo bash install-server.sh join {P['SERVER1_IP']}"
    )

    doc.add_heading("4.4 Agent Node Join (AGENT-1, AGENT-2)", level=2)
    add_code_block(
        doc,
        f"# SSH to each agent node\n"
        f"ssh root@{P['AGENT1_IP']}\n\n"
        f"# Run the agent install script\n"
        f"sudo bash install-agent.sh {P['SERVER1_IP']} <NODE_TOKEN>\n\n"
        f"# Repeat for AGENT-2\n"
        f"ssh root@{P['AGENT2_IP']}\n"
        f"sudo bash install-agent.sh {P['SERVER1_IP']} <NODE_TOKEN>"
    )

    doc.add_heading("4.5 Cluster Validation", level=2)
    doc.add_paragraph(
        "After all nodes have joined, execute the following verification on any server node:"
    )
    add_code_block(
        doc,
        f"# Configure kubectl access\n"
        f"export PATH=\"/var/lib/rancher/rke2/bin:$PATH\"\n"
        f"export KUBECONFIG=/etc/rancher/rke2/rke2.yaml\n\n"
        f"# Permanent configuration (add to ~/.bashrc)\n"
        f"echo 'export PATH=\"/var/lib/rancher/rke2/bin:$PATH\"' >> ~/.bashrc\n"
        f"echo 'export KUBECONFIG=/etc/rancher/rke2/rke2.yaml' >> ~/.bashrc\n\n"
        f"# Verify all nodes are Ready\n"
        f"kubectl get nodes -o wide\n\n"
        f"# Verify system pods\n"
        f"kubectl get pods -n kube-system\n\n"
        f"# Verify etcd cluster health\n"
        f"kubectl -n kube-system exec etcd-{P['SERVER1_NAME']} -- etcdctl \\\n"
        f"  --endpoints=https://127.0.0.1:2379 \\\n"
        f"  --cacert=/var/lib/rancher/rke2/server/tls/etcd/server-ca.crt \\\n"
        f"  --cert=/var/lib/rancher/rke2/server/tls/etcd/server-client.crt \\\n"
        f"  --key=/var/lib/rancher/rke2/server/tls/etcd/server-client.key \\\n"
        f"  endpoint health --cluster\n\n"
        f"# Verify etcd member list\n"
        f"kubectl -n kube-system exec etcd-{P['SERVER1_NAME']} -- etcdctl \\\n"
        f"  --endpoints=https://127.0.0.1:2379 \\\n"
        f"  --cacert=/var/lib/rancher/rke2/server/tls/etcd/server-ca.crt \\\n"
        f"  --cert=/var/lib/rancher/rke2/server/tls/etcd/server-client.crt \\\n"
        f"  --key=/var/lib/rancher/rke2/server/tls/etcd/server-client.key \\\n"
        f"  member list"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 5. LONGHORN DEPLOYMENT
    # ════════════════════════════════════════════════════════
    doc.add_heading("5. Longhorn Storage Deployment", level=1)

    doc.add_heading("5.1 Architecture Overview", level=2)
    doc.add_paragraph(
        "Longhorn is deployed as a distributed block storage system across the 3 server "
        "nodes, utilizing their 2TB disks for data replication. Each PersistentVolume is "
        "replicated 3 times (one replica per server node), providing fault tolerance "
        "against 2 simultaneous disk/node failures."
    )

    add_table_with_header(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["Default replica count", "3", "Replicas per volume (spread across 3 nodes)"],
            ["Data locality", "best-effort", "Prefer replica co-located with pod"],
            ["Filesystem", "ext4", "Volume filesystem type"],
            ["Reclaim policy", "Delete", "PVC deletion removes volume"],
            ["Stale replica timeout", "2880 min (48h)", "Grace period before marking stale"],
            ["Storage over-provisioning", "100%", "No over-provisioning"],
            ["Minimum available", "10%", "Reserved headroom per disk"],
            ["Backup target", f"s3://{P['S3_BUCKET']}@{P['S3_REGION']}", "S3-based backup repository"],
        ],
    )

    doc.add_heading("5.2 Disk Preparation", level=2)
    doc.add_paragraph(
        "On each server node, the 2TB disk must be formatted and mounted before Longhorn "
        "can use it. If using raw block devices, Longhorn can manage the filesystem."
    )
    add_code_block(
        doc,
        f"# On each server node (SERVER-1, SERVER-2, SERVER-3):\n\n"
        f"# Identify the 2TB disk\n"
        f"lsblk\n"
        f"fdisk -l | grep \"2 TB\"\n\n"
        f"# Option A: Raw disk (recommended — Longhorn manages filesystem)\n"
        f"# No preparation needed. Longhorn will format and manage the disk.\n"
        f"# Disk path example: /dev/sdb\n\n"
        f"# Option B: Pre-formatted disk\n"
        f"mkfs.ext4 /dev/sdb\n"
        f"mkdir -p /var/lib/longhorn\n"
        f"mount /dev/sdb /var/lib/longhorn\n"
        f"echo '/dev/sdb /var/lib/longhorn ext4 defaults 0 2' >> /etc/fstab"
    )

    doc.add_heading("5.3 Node Labeling", level=2)
    doc.add_paragraph(
        "Server nodes must be labeled to enable Longhorn default disk creation. "
        "Execute from any server node with kubectl access:"
    )
    add_code_block(
        doc,
        f"# Label server nodes for Longhorn default disk\n"
        f"kubectl label node {P['SERVER1_NAME']} node.longhorn.io/create-default-disk=true\n"
        f"kubectl label node {P['SERVER2_NAME']} node.longhorn.io/create-default-disk=true\n"
        f"kubectl label node {P['SERVER3_NAME']} node.longhorn.io/create-default-disk=true\n\n"
        f"# Optionally label agent nodes if they have dedicated storage\n"
        f"# kubectl label node {P['AGENT1_NAME']} node.longhorn.io/create-default-disk=true\n"
        f"# kubectl label node {P['AGENT2_NAME']} node.longhorn.io/create-default-disk=true"
    )

    doc.add_heading("5.4 Helm Installation", level=2)
    add_code_block(
        doc,
        f"# Add Longhorn Helm repository\n"
        f"helm repo add longhorn https://charts.longhorn.io\n"
        f"helm repo update\n\n"
        f"# Install Longhorn with custom values\n"
        f"helm upgrade --install longhorn longhorn/longhorn \\\n"
        f"  --namespace longhorn-system \\\n"
        f"  --create-namespace \\\n"
        f"  --version {P['LONGHORN_VERSION']} \\\n"
        f"  -f longhorn-values.yaml \\\n"
        f"  --wait --timeout 10m"
    )

    doc.add_heading("5.5 Helm Values (longhorn-values.yaml)", level=2)
    doc.add_paragraph("The following Helm values configure Longhorn for this deployment:")
    add_code_block(
        doc,
        f"defaultSettings:\n"
        f"  defaultReplicaCount: 3\n"
        f"  staleReplicaTimeout: 2880\n"
        f"  defaultDataLocality: best-effort\n"
        f"  createDefaultDiskLabeledNodes: true\n"
        f"  storageOverProvisioningPercentage: 100\n"
        f"  storageMinimalAvailablePercentage: 10\n"
        f"  autoUpgradeEngine: true\n"
        f"  replicaSoftAntiAffinity: true\n"
        f"  replicaZoneSoftAntiAffinity: true\n"
        f"  guaranteedInstanceManagerCpu: 0\n\n"
        f"persistence:\n"
        f"  defaultClass: true\n"
        f"  defaultClassReplicaCount: 3\n"
        f"  reclaimPolicy: Delete\n"
        f"  defaultFsType: ext4\n\n"
        f"service:\n"
        f"  ui:\n"
        f"    type: NodePort\n"
        f"    nodePort: 31080\n\n"
        f"longhornManager:\n"
        f"  tolerations:\n"
        f"    - key: \"node-role.kubernetes.io/control-plane\"\n"
        f"      operator: \"Exists\"\n"
        f"      effect: \"NoSchedule\"\n"
        f"  resources:\n"
        f"    requests:\n"
        f"      cpu: 100m\n"
        f"      memory: 128Mi\n"
        f"    limits:\n"
        f"      cpu: 500m\n"
        f"      memory: 512Mi"
    )

    doc.add_heading("5.6 Post-Installation Verification", level=2)
    add_code_block(
        doc,
        f"# Verify Longhorn pods\n"
        f"kubectl -n longhorn-system get pods -o wide\n\n"
        f"# Verify Longhorn nodes and disks\n"
        f"kubectl -n longhorn-system get nodes -o wide\n\n"
        f"# Verify StorageClass\n"
        f"kubectl get storageclass\n"
        f"# Expected: longhorn (default) → driver.longhorn.io\n\n"
        f"# Access Longhorn UI\n"
        f"# Option A: NodePort → http://<ANY_NODE_IP>:31080\n"
        f"# Option B: Port-forward → kubectl -n longhorn-system port-forward svc/longhorn-ui 8000:80"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 6. STORAGE CONFIGURATION
    # ════════════════════════════════════════════════════════
    doc.add_heading("6. Storage Configuration", level=1)

    doc.add_heading("6.1 StorageClass Definitions", level=2)
    doc.add_paragraph(
        "The following StorageClasses are deployed to provide differentiated "
        "storage tiers based on redundancy and performance requirements."
    )

    doc.add_heading("6.1.1 Default StorageClass (3 Replicas)", level=2)
    add_code_block(
        doc,
        f"apiVersion: storage.k8s.io/v1\n"
        f"kind: StorageClass\n"
        f"metadata:\n"
        f"  name: longhorn\n"
        f"  annotations:\n"
        f"    storageclass.kubernetes.io/is-default-class: \"true\"\n"
        f"provisioner: driver.longhorn.io\n"
        f"allowVolumeExpansion: true\n"
        f"reclaimPolicy: Delete\n"
        f"volumeBindingMode: WaitForFirstConsumer\n"
        f"parameters:\n"
        f"  numberOfReplicas: \"3\"\n"
        f"  staleReplicaTimeout: \"2880\"\n"
        f"  fromBackup: \"\"\n"
        f"  fsType: \"ext4\"\n"
        f"  dataLocality: \"best-effort\""
    )

    doc.add_heading("6.1.2 High-Performance StorageClass (Strict Local)", level=2)
    add_code_block(
        doc,
        f"apiVersion: storage.k8s.io/v1\n"
        f"kind: StorageClass\n"
        f"metadata:\n"
        f"  name: longhorn-high-perf\n"
        f"provisioner: driver.longhorn.io\n"
        f"allowVolumeExpansion: true\n"
        f"reclaimPolicy: Delete\n"
        f"volumeBindingMode: WaitForFirstConsumer\n"
        f"parameters:\n"
        f"  numberOfReplicas: \"3\"\n"
        f"  staleReplicaTimeout: \"2880\"\n"
        f"  fsType: \"ext4\"\n"
        f"  dataLocality: \"strict-local\""
    )

    doc.add_heading("6.1.3 Minimal StorageClass (2 Replicas)", level=2)
    add_code_block(
        doc,
        f"apiVersion: storage.k8s.io/v1\n"
        f"kind: StorageClass\n"
        f"metadata:\n"
        f"  name: longhorn-minimal\n"
        f"provisioner: driver.longhorn.io\n"
        f"allowVolumeExpansion: true\n"
        f"reclaimPolicy: Delete\n"
        f"volumeBindingMode: WaitForFirstConsumer\n"
        f"parameters:\n"
        f"  numberOfReplicas: \"2\"\n"
        f"  staleReplicaTimeout: \"2880\"\n"
        f"  fsType: \"ext4\"\n"
        f"  dataLocality: \"best-effort\""
    )

    doc.add_heading("6.2 PersistentVolumeClaim Examples", level=2)

    doc.add_heading("6.2.1 Standard PVC", level=2)
    add_code_block(
        doc,
        f"apiVersion: v1\n"
        f"kind: PersistentVolumeClaim\n"
        f"metadata:\n"
        f"  name: app-data-pvc\n"
        f"  namespace: default\n"
        f"spec:\n"
        f"  storageClassName: longhorn\n"
        f"  accessModes:\n"
        f"    - ReadWriteOnce\n"
        f"  resources:\n"
        f"    requests:\n"
        f"      storage: 50Gi"
    )

    doc.add_heading("6.2.2 StatefulSet with VolumeClaimTemplate", level=2)
    add_code_block(
        doc,
        f"apiVersion: apps/v1\n"
        f"kind: StatefulSet\n"
        f"metadata:\n"
        f"  name: mysql\n"
        f"spec:\n"
        f"  serviceName: mysql\n"
        f"  replicas: 1\n"
        f"  selector:\n"
        f"    matchLabels:\n"
        f"      app: mysql\n"
        f"  template:\n"
        f"    metadata:\n"
        f"      labels:\n"
        f"        app: mysql\n"
        f"    spec:\n"
        f"      containers:\n"
        f"        - name: mysql\n"
        f"          image: mysql:8.0\n"
        f"          env:\n"
        f"            - name: MYSQL_ROOT_PASSWORD\n"
        f"              valueFrom:\n"
        f"                secretKeyRef:\n"
        f"                  name: mysql-secret\n"
        f"                  key: password\n"
        f"          ports:\n"
        f"            - containerPort: 3306\n"
        f"          volumeMounts:\n"
        f"            - name: data\n"
        f"              mountPath: /var/lib/mysql\n"
        f"      tolerations:\n"
        f"        - key: \"node-role.kubernetes.io/control-plane\"\n"
        f"          operator: \"Exists\"\n"
        f"          effect: \"NoSchedule\"\n"
        f"  volumeClaimTemplates:\n"
        f"    - metadata:\n"
        f"        name: data\n"
        f"      spec:\n"
        f"        storageClassName: longhorn\n"
        f"        accessModes: [\"ReadWriteOnce\"]\n"
        f"        resources:\n"
        f"          requests:\n"
        f"            storage: 100Gi"
    )

    doc.add_heading("6.3 Volume Expansion", level=2)
    doc.add_paragraph(
        "Longhorn supports online volume expansion. To expand an existing PVC:"
    )
    add_code_block(
        doc,
        f"# Edit the PVC to increase size\n"
        f"kubectl edit pvc app-data-pvc\n"
        f"# Change: spec.resources.requests.storage: 50Gi → 100Gi\n\n"
        f"# Verify expansion\n"
        f"kubectl get pvc app-data-pvc\n"
        f"kubectl -n longhorn-system get volumes"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 7. OPERATIONS
    # ════════════════════════════════════════════════════════
    doc.add_heading("7. Operational Procedures", level=1)

    doc.add_heading("7.1 Cluster Health Monitoring", level=2)
    add_code_block(
        doc,
        f"# Node status\n"
        f"kubectl get nodes -o wide\n\n"
        f"# Resource utilization (requires metrics-server)\n"
        f"kubectl top nodes\n"
        f"kubectl top pods -A --sort-by=memory\n\n"
        f"# System pod health\n"
        f"kubectl get pods -n kube-system -o wide\n\n"
        f"# etcd endpoint health\n"
        f"kubectl -n kube-system exec etcd-{P['SERVER1_NAME']} -- etcdctl \\\n"
        f"  --endpoints=https://127.0.0.1:2379 \\\n"
        f"  --cacert=/var/lib/rancher/rke2/server/tls/etcd/server-ca.crt \\\n"
        f"  --cert=/var/lib/rancher/rke2/server/tls/etcd/server-client.crt \\\n"
        f"  --key=/var/lib/rancher/rke2/server/tls/etcd/server-client.key \\\n"
        f"  endpoint health\n\n"
        f"# etcd database size\n"
        f"kubectl -n kube-system exec etcd-{P['SERVER1_NAME']} -- etcdctl \\\n"
        f"  --endpoints=https://127.0.0.1:2379 \\\n"
        f"  --cacert=/var/lib/rancher/rke2/server/tls/etcd/server-ca.crt \\\n"
        f"  --cert=/var/lib/rancher/rke2/server/tls/etcd/server-client.crt \\\n"
        f"  --key=/var/lib/rancher/rke2/server/tls/etcd/server-client.key \\\n"
        f"  endpoint status --write-out=table"
    )

    doc.add_heading("7.2 Longhorn Monitoring", level=2)
    add_code_block(
        doc,
        f"# Volume status\n"
        f"kubectl -n longhorn-system get volumes\n\n"
        f"# Replica status\n"
        f"kubectl -n longhorn-system get replicas\n\n"
        f"# Node disk usage\n"
        f"kubectl -n longhorn-system get nodes -o custom-columns=\\\n"
        f"NODE:.metadata.name,\\\n"
        f"DISK:.status.diskStatus\n\n"
        f"# Longhorn manager logs\n"
        f"kubectl -n longhorn-system logs -l app=longhorn-manager --tail=100\n\n"
        f"# Engine logs for a specific volume\n"
        f"kubectl -n longhorn-system logs -l longhorn.io/volume=<VOLUME_NAME>"
    )

    doc.add_heading("7.3 RKE2 Service Management", level=2)
    add_table_with_header(
        doc,
        ["Operation", "Server Node Command", "Agent Node Command"],
        [
            ["Start", "systemctl start rke2-server", "systemctl start rke2-agent"],
            ["Stop", "systemctl stop rke2-server", "systemctl stop rke2-agent"],
            ["Restart", "systemctl restart rke2-server", "systemctl restart rke2-agent"],
            ["Status", "systemctl status rke2-server", "systemctl status rke2-agent"],
            ["Logs", "journalctl -u rke2-server -f", "journalctl -u rke2-agent -f"],
            ["Enable on boot", "systemctl enable rke2-server", "systemctl enable rke2-agent"],
        ],
    )

    add_note(
        doc,
        "When restarting server nodes, restart one at a time and wait for it to become "
        "Ready before restarting the next. Simultaneous restarts of 2+ server nodes will "
        "cause etcd quorum loss.",
        "WARNING"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 8. SCALING
    # ════════════════════════════════════════════════════════
    doc.add_heading("8. Scaling Procedures", level=1)

    doc.add_heading("8.1 Horizontal Node Scaling", level=2)

    doc.add_heading("8.1.1 Adding a Server Node", level=2)
    doc.add_paragraph(
        "To expand the control plane from 3 to 5 server nodes (improves etcd fault "
        "tolerance from 1 to 2 node failures):"
    )
    add_code_block(
        doc,
        f"# On the new server node:\n\n"
        f"# 1. Prepare OS (disable swap, firewalld, SELinux, sync NTP)\n"
        f"# 2. Run server join script\n"
        f"sudo bash install-server.sh join {P['SERVER1_IP']}\n\n"
        f"# 3. Verify on any existing server\n"
        f"kubectl get nodes\n\n"
        f"# 4. Label for Longhorn (if node has 2TB+ disk)\n"
        f"kubectl label node <NEW_NODE> node.longhorn.io/create-default-disk=true\n\n"
        f"# 5. Update load balancer configuration\n"
        f"# Add the new node IP to the backend pool on port 6443"
    )

    add_note(
        doc,
        "Always add server nodes in odd numbers (3 → 5 → 7) to maintain etcd quorum. "
        "Adding an even number provides no additional fault tolerance.",
        "IMPORTANT"
    )

    doc.add_heading("8.1.2 Adding an Agent Node", level=2)
    add_code_block(
        doc,
        f"# On the new agent node:\n\n"
        f"# 1. Prepare OS (disable swap, firewalld, SELinux, sync NTP)\n"
        f"# 2. Run agent join script\n"
        f"sudo bash install-agent.sh {P['SERVER1_IP']} <NODE_TOKEN>\n\n"
        f"# 3. Verify\n"
        f"kubectl get nodes\n\n"
        f"# 4. Label for Longhorn (optional, if node has storage)\n"
        f"kubectl label node <NEW_NODE> node.longhorn.io/create-default-disk=true"
    )

    doc.add_heading("8.1.3 Removing a Node", level=2)
    add_code_block(
        doc,
        f"# 1. Cordon the node (prevent new pod scheduling)\n"
        f"kubectl cordon <NODE_NAME>\n\n"
        f"# 2. Drain the node (evict existing pods)\n"
        f"kubectl drain <NODE_NAME> --ignore-daemonsets --delete-emptydir-data\n\n"
        "# 3. If Longhorn node: disable disk scheduling\n"
        "kubectl -n longhorn-system patch node <NODE_NAME> -p \\\n"
        '  \'{"spec":{"allowScheduling":false}}\'\n\n'
        f"# 4. Wait for Longhorn replicas to migrate\n"
        f"kubectl -n longhorn-system get replicas | grep <NODE_NAME>\n\n"
        f"# 5. Remove from cluster\n"
        f"kubectl delete node <NODE_NAME>\n\n"
        f"# 6. Cleanup on the removed node\n"
        f"sudo systemctl stop rke2-server  # or rke2-agent\n"
        f"sudo systemctl disable rke2-server  # or rke2-agent\n"
        f"sudo rm -rf /var/lib/rancher/rke2 /etc/rancher/rke2"
    )

    doc.add_heading("8.2 Vertical Storage Scaling", level=2)

    doc.add_heading("8.2.1 Adding Disk to Existing Node", level=2)
    add_code_block(
        doc,
        f"# 1. Attach new disk (e.g., /dev/sdc) to the VM/physical server\n\n"
        f"# 2. On the node, verify disk visibility\n"
        f"lsblk\n\n"
        f"# 3. Via Longhorn UI: Node → Edit Disks → Add Disk\n"
        f"#    Path: /dev/sdc\n"
        f"#    Allow Scheduling: true\n"
        f"#    Storage Reserved: 0 (or specify reserved amount)\n\n"
        f"# Or via kubectl:\n"
        f"kubectl -n longhorn-system edit node <NODE_NAME>"
    )

    doc.add_heading("8.2.2 Expanding Existing Disk", level=2)
    add_code_block(
        doc,
        f"# 1. Expand disk at hypervisor/storage level (e.g., 2TB → 4TB)\n\n"
        f"# 2. Rescan SCSI bus (if needed)\n"
        f"echo 1 > /sys/class/block/sdb/device/rescan\n\n"
        f"# 3. Expand partition\n"
        f"growpart /dev/sdb 1\n\n"
        f"# 4. Resize filesystem\n"
        f"resize2fs /dev/sdb1  # ext4\n"
        f"# OR\n"
        f"xfs_growfs /mount/point  # XFS\n\n"
        f"# 5. Longhorn auto-detects the new capacity\n"
        f"kubectl -n longhorn-system get nodes -o yaml | grep storageMaximum"
    )

    doc.add_heading("8.3 Increasing Volume Replica Count", level=2)
    add_code_block(
        doc,
        f"# Create a high-availability StorageClass with 5 replicas\n"
        f"cat <<EOF | kubectl apply -f -\n"
        f"apiVersion: storage.k8s.io/v1\n"
        f"kind: StorageClass\n"
        f"metadata:\n"
        f"  name: longhorn-ha\n"
        f"provisioner: driver.longhorn.io\n"
        f"allowVolumeExpansion: true\n"
        f"reclaimPolicy: Retain\n"
        f"volumeBindingMode: WaitForFirstConsumer\n"
        f"parameters:\n"
        f"  numberOfReplicas: \"5\"\n"
        f"  staleReplicaTimeout: \"2880\"\n"
        f"  fsType: \"ext4\"\n"
        f"  dataLocality: \"best-effort\"\n"
        f"EOF\n\n"
        f"# For existing volumes, edit via Longhorn UI or kubectl:\n"
        f"kubectl -n longhorn-system edit volume <VOLUME_NAME>\n"
        f"# Change: spec.numberOfReplicas: 3 → 5"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 9. BACKUP & DR
    # ════════════════════════════════════════════════════════
    doc.add_heading("9. Backup and Disaster Recovery", level=1)

    doc.add_heading("9.1 etcd Backup", level=2)
    doc.add_paragraph(
        "RKE2 automatically takes etcd snapshots every 6 hours. Snapshots are stored "
        "on each server node at /var/lib/rancher/rke2/server/db/snapshots/."
    )
    add_code_block(
        doc,
        f"# Manual snapshot\n"
        f"sudo rke2 etcd-snapshot save --name manual-backup-$(date +%Y%m%d)\n\n"
        f"# List snapshots\n"
        f"sudo rke2 etcd-snapshot ls\n\n"
        f"# Copy snapshot off-node\n"
        f"scp /var/lib/rancher/rke2/server/db/snapshots/<SNAPSHOT> backup@backup-server:/backups/"
    )

    doc.add_heading("9.2 etcd Restore", level=2)
    add_code_block(
        doc,
        f"# 1. Stop RKE2 on ALL server nodes\n"
        f"# On each server node:\n"
        f"sudo systemctl stop rke2-server\n\n"
        f"# 2. On one node, restore from snapshot\n"
        f"sudo rke2 server \\\n"
        f"  --cluster-reset \\\n"
        f"  --cluster-reset-restore-path=/var/lib/rancher/rke2/server/db/snapshots/<SNAPSHOT>\n\n"
        f"# 3. Start the restored node\n"
        f"sudo systemctl start rke2-server\n\n"
        f"# 4. On remaining nodes, remove etcd data and rejoin\n"
        f"sudo rm -rf /var/lib/rancher/rke2/server/db\n"
        f"sudo systemctl start rke2-server\n\n"
        f"# 5. Verify\n"
        f"kubectl get nodes"
    )

    doc.add_heading("9.3 Longhorn Volume Backup", level=2)
    add_code_block(
        doc,
        f"# Configure S3 backup target (via Longhorn UI → Settings → Backup Target)\n"
        f"backup-target: s3://{P['S3_BUCKET']}@{P['S3_REGION']}/\n"
        f"backup-target-credential-secret: longhorn-backup-secret\n\n"
        f"# Create backup (via UI: Volume → Create Backup)\n"
        f"# Or via kubectl:\n"
        f"cat <<EOF | kubectl apply -f -\n"
        f"apiVersion: longhorn.io/v1beta2\n"
        f"kind: Backup\n"
        f"metadata:\n"
        f"  name: manual-backup-$(date +%Y%m%d)\n"
        f"  namespace: longhorn-system\n"
        f"spec:\n"
        f"  snapshotName: <SNAPSHOT_NAME>\n"
        f"  volumeName: <VOLUME_NAME>\n"
        f"EOF"
    )

    doc.add_heading("9.4 Longhorn Volume Restore", level=2)
    add_code_block(
        doc,
        f"# Restore from backup (via Longhorn UI: Backup → Restore)\n"
        f"# Or via kubectl — create PVC from backup:\n"
        f"cat <<EOF | kubectl apply -f -\n"
        f"apiVersion: v1\n"
        f"kind: PersistentVolumeClaim\n"
        f"metadata:\n"
        f"  name: restored-pvc\n"
        f"spec:\n"
        f"  storageClassName: longhorn\n"
        f"  accessModes:\n"
        f"    - ReadWriteOnce\n"
        f"  resources:\n"
        f"    requests:\n"
        f"      storage: 50Gi\n"
        f"  dataSource:\n"
        f"    name: <BACKUP_NAME>\n"
        f"    kind: Backup\n"
        f"    apiGroup: longhorn.io\n"
        f"EOF"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 10. TROUBLESHOOTING
    # ════════════════════════════════════════════════════════
    doc.add_heading("10. Troubleshooting Guide", level=1)

    troubleshooting_items = [
        {
            "title": "10.1 Node in NotReady State",
            "symptoms": "kubectl get nodes shows NotReady status",
            "diagnosis": [
                "kubectl describe node <NODE> — check Conditions and Events",
                "SSH to node: systemctl status rke2-server (or rke2-agent)",
                "journalctl -u rke2-server -n 200 — check for errors",
            ],
            "resolution": [
                "If kubelet crashed: systemctl restart rke2-server",
                "If disk pressure: clean up images (crictl rmi --prune)",
                "If memory pressure: reduce pod density or add nodes",
                "If certificate expired: renew with rke2 certificate rotate",
            ],
        },
        {
            "title": "10.2 Pod Stuck in Pending",
            "symptoms": "Pod shows Pending status, no node assignment",
            "diagnosis": [
                "kubectl describe pod <POD> — check Events for scheduling failure",
                "Check node resources: kubectl top nodes",
                "Check taints: kubectl get nodes -o custom-columns=NODE:.metadata.name,TAINTS:.spec.taints",
            ],
            "resolution": [
                "Insufficient resources: scale nodes or reduce resource requests",
                "Taint mismatch: add toleration to pod spec",
                "PVC binding issue: check PVC status and Longhorn health",
            ],
        },
        {
            "title": "10.3 PVC Stuck in Pending",
            "symptoms": "PVC won't bind to PV, remains Pending",
            "diagnosis": [
                "kubectl describe pvc <PVC> — check Events",
                "kubectl -n longhorn-system get volumes — check Longhorn volumes",
                "kubectl -n longhorn-system get nodes — check disk availability",
            ],
            "resolution": [
                "If Longhorn not running: check longhorn-system pods",
                "If disk full: expand disk or delete unused volumes",
                "If StorageClass missing: apply storageclass.yaml",
            ],
        },
        {
            "title": "10.4 etcd Quorum Loss",
            "symptoms": "API server returns errors, cluster partially unresponsive",
            "diagnosis": [
                "kubectl -n kube-system get pods | grep etcd — check etcd pod status",
                "Check if 2+ server nodes are down",
            ],
            "resolution": [
                "Bring up failed server nodes: systemctl start rke2-server",
                "If nodes unrecoverable: restore etcd from snapshot (see Section 9.2)",
                "If single node: wait for auto-recovery or restart rke2-server",
            ],
        },
        {
            "title": "10.5 Longhorn Volume Degraded",
            "symptoms": "Volume shows Degraded state, fewer than expected replicas",
            "diagnosis": [
                "kubectl -n longhorn-system get volumes — check replica count",
                "kubectl -n longhorn-system get replicas — check individual replica state",
                "kubectl -n longhorn-system logs -l app=longhorn-manager — check for errors",
            ],
            "resolution": [
                "Auto-heal: Longhorn will rebuild replicas on healthy nodes automatically",
                "Manual trigger: via Longhorn UI → Volume → Salvage",
                "If node is permanently down: remove node, replicas rebuild elsewhere",
            ],
        },
        {
            "title": "10.6 Node Join Failure",
            "symptoms": "New node fails to join, error in logs",
            "diagnosis": [
                "journalctl -u rke2-server -f — check join errors",
                "Verify connectivity: ping <SERVER-1-IP>",
                "Verify token: cat /var/lib/rancher/rke2/server/node-token on bootstrap node",
            ],
            "resolution": [
                "Certificate error: verify token and server IP are correct",
                "Connection refused: check firewall, verify port 9345 is open",
                "Node already exists: clean up with rm -rf /var/lib/rancher/rke2 and retry",
            ],
        },
    ]

    for item in troubleshooting_items:
        doc.add_heading(item["title"], level=2)

        p = doc.add_paragraph()
        run = p.add_run("Symptoms: ")
        run.bold = True
        p.add_run(item["symptoms"])

        p = doc.add_paragraph()
        run = p.add_run("Diagnosis:")
        run.bold = True
        for d in item["diagnosis"]:
            add_code_block(doc, d)

        p = doc.add_paragraph()
        run = p.add_run("Resolution:")
        run.bold = True
        for r in item["resolution"]:
            doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("10.7 Diagnostic Commands Reference", level=2)
    add_code_block(
        doc,
        f"# Full cluster diagnostic dump\n"
        f"echo '=== NODES ===' && kubectl get nodes -o wide\n"
        f"echo '=== PODS ===' && kubectl get pods -A -o wide\n"
        f"echo '=== PVCs ===' && kubectl get pvc -A\n"
        f"echo '=== SERVICES ===' && kubectl get svc -A\n"
        f"echo '=== EVENTS ===' && kubectl get events --sort-by=.lastTimestamp -A | tail -30\n"
        f"echo '=== LONGHORN ===' && kubectl -n longhorn-system get volumes,replicas,nodes\n"
        f"echo '=== ETCD ===' && kubectl -n kube-system get pods | grep etcd"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # APPENDIX A: FILE INVENTORY
    # ════════════════════════════════════════════════════════
    doc.add_heading("Appendix A: Deployment File Inventory", level=1)
    add_table_with_header(
        doc,
        ["File", "Location", "Purpose"],
        [
            ["install-server.sh", "rke2-setup/scripts/", "Server node installation (bootstrap + join)"],
            ["install-agent.sh", "rke2-setup/scripts/", "Agent node installation"],
            ["install-longhorn.sh", "rke2-setup/scripts/", "Longhorn Helm deployment"],
            ["setup-complete.sh", "rke2-setup/scripts/", "Post-install manifest application"],
            ["longhorn-values.yaml", "rke2-setup/", "Longhorn Helm values"],
            ["storageclass.yaml", "rke2-setup/manifests/", "StorageClass definitions"],
            ["sample-workloads.yaml", "rke2-setup/manifests/", "Example PVC, Pod, StatefulSet"],
        ],
    )

    # ════════════════════════════════════════════════════════
    # APPENDIX B: PORT REFERENCE
    # ════════════════════════════════════════════════════════
    doc.add_heading("Appendix B: Complete Port Reference", level=1)
    add_table_with_header(
        doc,
        ["Port", "Protocol", "Component", "Direction", "Nodes"],
        [
            ["6443", "TCP", "kube-apiserver", "Inbound", "All"],
            ["9345", "TCP", "RKE2 supervisor", "Inbound", "Servers"],
            ["2379", "TCP", "etcd client", "Inbound", "Servers"],
            ["2380", "TCP", "etcd peer", "Inbound", "Servers"],
            ["10250", "TCP", "kubelet", "Inbound", "All"],
            ["10251", "TCP", "kube-scheduler", "Inbound", "Servers"],
            ["10252", "TCP", "kube-controller", "Inbound", "Servers"],
            ["4789", "UDP", "VXLAN (Canal)", "Bidirectional", "All"],
            ["8472", "UDP", "VXLAN (Canal)", "Bidirectional", "All"],
            ["9500", "TCP", "Longhorn manager", "Inbound", "All"],
            ["9501", "TCP", "Longhorn replica", "Inbound", "All"],
            ["9502", "TCP", "Longhorn engine", "Inbound", "All"],
            ["30000-32767", "TCP/UDP", "NodePort range", "Inbound", "All"],
            ["31080", "TCP", "Longhorn UI", "Inbound", "All"],
        ],
    )

    # ════════════════════════════════════════════════════════
    # APPENDIX C: DIRECTORY STRUCTURE
    # ════════════════════════════════════════════════════════
    doc.add_heading("Appendix C: RKE2 Directory Structure", level=1)
    add_code_block(
        doc,
        "/etc/rancher/rke2/\n"
        "├── config.yaml              # Cluster configuration\n"
        "├── rke2.yaml                # Kubeconfig (generated)\n"
        "└── certs/                   # TLS certificates\n\n"
        "/var/lib/rancher/rke2/\n"
        "├── agent/                   # kubelet data\n"
        "│   ├── containerd/          # Container runtime data\n"
        "│   └── etc/                 # Agent configuration\n"
        "├── server/\n"
        "│   ├── db/                  # etcd database\n"
        "│   │   └── snapshots/       # etcd snapshots\n"
        "│   ├── tls/                 # Server TLS certificates\n"
        "│   ├── manifests/           # Static pod manifests\n"
        "│   ├── cred/                # Admin kubeconfig\n"
        "│   └── logs/                # Audit and server logs\n"
        "└── bin/                     # kubectl, crictl, etc."
    )

    # ════════════════════════════════════════════════════════
    # SAVE
    # ════════════════════════════════════════════════════════
    output_path = "/root/opencode/rke2-setup/RKE2-Longhorn-Deployment-Guide.docx"
    doc.save(output_path)
    print(f"Document generated: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
